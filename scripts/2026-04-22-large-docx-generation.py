from docx import Document
from docx.enum.text import WD_BREAK
import random
import string

def generate_random_text(length=1000):
    words = []
    for _ in range(length):
        word_len = random.randint(3, 10)
        word = ''.join(random.choices(string.ascii_lowercase, k=word_len))
        words.append(word)
    return ' '.join(words)

def create_large_docx(filename, num_pages=1000):
    doc = Document()

    for page in range(1, num_pages + 1):
        header = f"=== PAGE {page} ==="
        doc.add_paragraph(header)

        text = generate_random_text(300)
        doc.add_paragraph(text)

        extra = f"\n[UNIQUE_ID: {page}_{random.randint(100000, 999999)}]"
        doc.add_paragraph(extra)

        if page != num_pages:
            doc.add_page_break()

        if page % 50 == 0:
            print(f"Сгенерировано страниц: {page}")

    doc.save(filename)
    print(f"Файл сохранен как {filename}")


if __name__ == "__main__":
    create_large_docx("huge_document.docx", num_pages=1000)